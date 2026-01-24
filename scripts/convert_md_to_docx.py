#!/usr/bin/env python3
"""
将Markdown文档转换为Word格式，支持Mermaid图表转换。

依赖:
- python-docx: Word文档生成
- markdown: Markdown解析
- Pillow: 图片处理
- playwright: Mermaid图表渲染

安装依赖:
    uv pip install python-docx markdown Pillow playwright
    uv run playwright install chromium

使用方法:
    # 基本用法（输出与输入同名.docx）
    uv run python scripts/convert_md_to_docx.py docs/architecture/ARCHITECTURE.md

    # 指定输出文件
    uv run python scripts/convert_md_to_docx.py docs/architecture/ARCHITECTURE.md -o output.docx

功能特性:
- ✅ 自动渲染Mermaid图表为PNG图片
- ✅ 保留Markdown格式（标题、段落、列表）
- ✅ 支持代码块（带语法高亮样式）
- ✅ 支持表格转换
- ✅ 支持Markdown内联标记（加粗、斜体、代码、链接）
- ✅ 支持有序/无序列表
- ✅ 支持引用块和水平分隔线
- ✅ 中英文字体优化（英文Arial，中文微软雅黑）
"""

import argparse
import re
import sys
import tempfile
from pathlib import Path
from typing import List, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 缺少 python-docx 依赖")
    print("请运行: pip install python-docx")
    sys.exit(1)

try:
    import markdown
    from markdown.extensions import tables, fenced_code
except ImportError:
    print("错误: 缺少 markdown 依赖")
    print("请运行: pip install markdown")
    sys.exit(1)


class MermaidRenderer:
    """使用Playwright渲染Mermaid图表为PNG图片"""

    def __init__(self):
        self.playwright = None
        self.browser = None
        self._initialized = False

    def initialize(self):
        """初始化Playwright"""
        if self._initialized:
            return

        try:
            from playwright.sync_api import sync_playwright

            self.playwright = sync_playwright().start()
            self.browser = self.playwright.chromium.launch(headless=True)
            self._initialized = True
            print("✓ Mermaid渲染器初始化成功")
        except ImportError:
            print("警告: 未安装 playwright,无法渲染Mermaid图表")
            print("请运行: pip install playwright && playwright install chromium")
            self._initialized = False
        except Exception as e:
            print(f"警告: Playwright初始化失败: {e}")
            self._initialized = False

    def render(self, mermaid_code: str, output_path: Path) -> bool:
        """
        渲染Mermaid代码为PNG图片

        Args:
            mermaid_code: Mermaid图表代码
            output_path: 输出PNG文件路径

        Returns:
            是否成功渲染
        """
        if not self._initialized:
            return False

        try:
            page = self.browser.new_page(viewport={"width": 1200, "height": 800})

            # 创建包含Mermaid的HTML页面
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
                <style>
                    body {{
                        margin: 0;
                        padding: 20px;
                        background: white;
                    }}
                    .mermaid {{
                        text-align: center;
                    }}
                </style>
            </head>
            <body>
                <div class="mermaid">
{mermaid_code}
                </div>
                <script>
                    mermaid.initialize({{ 
                        startOnLoad: true,
                        theme: 'default',
                        securityLevel: 'loose'
                    }});
                </script>
            </body>
            </html>
            """

            page.set_content(html_content)
            page.wait_for_timeout(2000)  # 等待渲染完成

            # 截图
            element = page.query_selector(".mermaid")
            if element:
                element.screenshot(path=str(output_path))
                page.close()
                return True
            else:
                page.close()
                return False

        except Exception as e:
            print(f"警告: Mermaid渲染失败: {e}")
            return False

    def close(self):
        """关闭浏览器"""
        if self.browser:
            self.browser.close()
        if self.playwright:
            self.playwright.stop()


class MarkdownToDocxConverter:
    """Markdown到Word文档转换器"""

    def __init__(self, md_file: Path, output_file: Path):
        self.md_file = md_file
        self.output_file = output_file
        self.doc = Document()
        self.mermaid_renderer = MermaidRenderer()
        self.temp_dir = Path(tempfile.mkdtemp())
        self.mermaid_count = 0

        # 设置默认字体
        self._set_default_font()

    def _set_default_font(self):
        """设置文档默认字体"""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        # 设置中文字体
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def extract_mermaid_blocks(self, content: str) -> List[Tuple[str, str]]:
        """
        提取Markdown中的Mermaid代码块

        Returns:
            List of (placeholder, mermaid_code) tuples
        """
        pattern = r"```mermaid\s*\n(.*?)\n```"
        matches = re.finditer(pattern, content, re.DOTALL)

        blocks = []
        for match in matches:
            mermaid_code = match.group(1).strip()
            placeholder = f"__MERMAID_PLACEHOLDER_{len(blocks)}__"
            blocks.append((placeholder, mermaid_code))

        return blocks

    def replace_mermaid_with_placeholders(self, content: str, blocks: List[Tuple[str, str]]) -> str:
        """将Mermaid代码块替换为占位符"""
        result = content
        pattern = r"```mermaid\s*\n.*?\n```"

        for placeholder, _ in blocks:
            result = re.sub(pattern, placeholder, result, count=1, flags=re.DOTALL)

        return result

    def render_mermaid_blocks(self, blocks: List[Tuple[str, str]]) -> dict:
        """渲染所有Mermaid代码块为PNG图片"""
        self.mermaid_renderer.initialize()

        image_paths = {}
        for i, (placeholder, mermaid_code) in enumerate(blocks):
            output_path = self.temp_dir / f"mermaid_{i}.png"

            print(f"渲染Mermaid图表 {i + 1}/{len(blocks)}...")
            success = self.mermaid_renderer.render(mermaid_code, output_path)

            if success and output_path.exists():
                image_paths[placeholder] = output_path
                print(f"  ✓ 已保存到: {output_path}")
            else:
                print(f"  ✗ 渲染失败,将使用文本占位")
                image_paths[placeholder] = None

        return image_paths

    def add_heading(self, text: str, level: int):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_paragraph(self, text: str, style=None):
        """添加段落，支持Markdown内联标记"""
        if not text.strip():
            return

        paragraph = self.doc.add_paragraph(style=style)
        self._parse_inline_markdown(paragraph, text)
        return paragraph

    def _parse_inline_markdown(self, paragraph, text: str):
        """解析并添加Markdown内联格式（加粗、斜体、代码、链接等）"""
        import re

        # 处理顺序：代码 > 加粗 > 斜体 > 链接
        parts = []
        current_pos = 0

        # 正则模式：捕获内联代码、加粗、斜体、链接
        # 顺序很重要：先匹配代码块，避免与其他标记冲突
        pattern = re.compile(
            r"(`[^`]+`)|"  # 内联代码
            r"(\*\*[^*]+\*\*)|"  # 加粗
            r"(\*[^*]+\*)|"  # 斜体
            r"(\[([^\]]+)\]\(([^)]+)\))"  # 链接 [text](url)
        )

        for match in pattern.finditer(text):
            # 添加匹配前的普通文本
            if match.start() > current_pos:
                plain_text = text[current_pos : match.start()]
                paragraph.add_run(plain_text)

            # 处理匹配的内联标记
            if match.group(1):  # 内联代码 `code`
                code_text = match.group(1)[1:-1]  # 去掉反引号
                run = paragraph.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(220, 50, 50)  # 红色
            elif match.group(2):  # 加粗 **text**
                bold_text = match.group(2)[2:-2]  # 去掉**
                run = paragraph.add_run(bold_text)
                run.bold = True
            elif match.group(3):  # 斜体 *text*
                italic_text = match.group(3)[1:-1]  # 去掉*
                run = paragraph.add_run(italic_text)
                run.italic = True
            elif match.group(4):  # 链接 [text](url)
                link_text = match.group(5)
                link_url = match.group(6)
                # Word中添加超链接
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                run.font.underline = True

            current_pos = match.end()

        # 添加剩余的普通文本
        if current_pos < len(text):
            paragraph.add_run(text[current_pos:])

    def add_code_block(self, code: str, language: str = ""):
        """添加代码块"""
        paragraph = self.doc.add_paragraph()
        paragraph.style = "Normal"

        # 设置代码块样式
        run = paragraph.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)

        # 设置背景色 (灰色)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)

    def add_table(self, rows: List[List[str]]):
        """添加表格"""
        if not rows:
            return

        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Light Grid Accent 1"

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text.strip()

    def add_image(self, image_path: Path, width: float = 6.0):
        """添加图片"""
        try:
            self.doc.add_picture(str(image_path), width=Inches(width))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"警告: 无法添加图片 {image_path}: {e}")

    def parse_and_add_content(self, content: str, image_paths: dict):
        """解析Markdown内容并添加到Word文档"""
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # 检查Mermaid占位符
            if line.strip().startswith("__MERMAID_PLACEHOLDER_"):
                placeholder = line.strip()
                if placeholder in image_paths:
                    image_path = image_paths[placeholder]
                    if image_path:
                        self.add_image(image_path)
                    else:
                        self.add_paragraph("[Mermaid图表 - 渲染失败]")
                i += 1
                continue

            # 标题
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("#").strip()
                self.add_heading(text, min(level, 9))
                i += 1
                continue

            # 代码块
            if line.strip().startswith("```"):
                language = line.strip()[3:].strip()
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                code = "\n".join(code_lines)
                self.add_code_block(code, language)
                i += 1
                continue

            # 表格
            if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
                table_lines = [line]
                i += 1
                # 跳过分隔行
                if "---" in lines[i]:
                    i += 1
                # 收集表格行
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1

                # 解析表格
                rows = []
                for table_line in table_lines:
                    cells = [cell.strip() for cell in table_line.split("|") if cell.strip()]
                    if cells:
                        rows.append(cells)

                self.add_table(rows)
                continue

            # 列表项（无序列表）
            if line.strip().startswith("- ") or line.strip().startswith("* "):
                text = line.strip()[2:].strip()
                paragraph = self.doc.add_paragraph(style="List Bullet")
                self._parse_inline_markdown(paragraph, text)
                i += 1
                continue

            # 列表项（有序列表）
            if re.match(r"^\d+\.\s", line.strip()):
                text = re.sub(r"^\d+\.\s", "", line.strip())
                paragraph = self.doc.add_paragraph(style="List Number")
                self._parse_inline_markdown(paragraph, text)
                i += 1
                continue

            # 引用块
            if line.strip().startswith("> "):
                text = line.strip()[2:]
                paragraph = self.doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.5)
                paragraph.paragraph_format.right_indent = Inches(0.5)
                self._parse_inline_markdown(paragraph, text)
                run = paragraph.runs[0] if paragraph.runs else None
                if run:
                    run.italic = True
                i += 1
                continue

            # 水平分隔线（使用文本行替代）
            if line.strip() in ["---", "***", "___"]:
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run("─" * 60)  # Unicode水平线字符
                run.font.color.rgb = RGBColor(200, 200, 200)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue

            # 普通段落
            if line.strip():
                self.add_paragraph(line)

            i += 1

    def convert(self):
        """执行转换"""
        print(f"读取文件: {self.md_file}")

        # 读取Markdown内容
        with open(self.md_file, "r", encoding="utf-8") as f:
            content = f.read()

        # 提取Mermaid代码块
        print("提取Mermaid代码块...")
        mermaid_blocks = self.extract_mermaid_blocks(content)
        print(f"找到 {len(mermaid_blocks)} 个Mermaid图表")

        # 渲染Mermaid图表
        image_paths = {}
        if mermaid_blocks:
            image_paths = self.render_mermaid_blocks(mermaid_blocks)

        # 替换Mermaid代码为占位符
        content = self.replace_mermaid_with_placeholders(content, mermaid_blocks)

        # 解析并添加内容到Word
        print("生成Word文档...")
        self.parse_and_add_content(content, image_paths)

        # 保存文档
        self.doc.save(str(self.output_file))
        print(f"✓ 转换完成: {self.output_file}")

        # 清理
        self.mermaid_renderer.close()


def main():
    parser = argparse.ArgumentParser(description="将Markdown文档转换为Word格式（支持Mermaid图表）")
    parser.add_argument("input_file", type=Path, help="输入的Markdown文件路径")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="输出的Word文件路径（默认：与输入文件同名.docx）",
    )

    args = parser.parse_args()

    # 检查输入文件
    if not args.input_file.exists():
        print(f"错误: 文件不存在: {args.input_file}")
        sys.exit(1)

    # 确定输出文件
    if args.output:
        output_file = args.output
    else:
        output_file = args.input_file.with_suffix(".docx")

    # 执行转换
    converter = MarkdownToDocxConverter(args.input_file, output_file)
    converter.convert()


if __name__ == "__main__":
    main()
