from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:
    from docx.oxml.numbering import CT_Numbering
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "MemStack-Desktop-Agent-Client-PRD.docx"

INK = RGBColor(0x19, 0x25, 0x33)
MUTED = RGBColor(0x67, 0x73, 0x86)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D7DEE7"
BASE_FONT = "Arial Unicode MS"
EAST_ASIA_FONT = BASE_FONT


def set_run_font(
    run: Run,
    *,
    name: str = BASE_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell: _Cell,
    *,
    top: int = 80,
    start: int = 120,
    bottom: int = 80,
    end: int = 120,
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, result, fld_end])


def set_paragraph_border_bottom(
    paragraph: Paragraph,
    color: str = BORDER,
    size: int = 8,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc: Document) -> None:  # noqa: PLR0915
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in [style.name for style in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["Code Block"]
    code_style.font.name = "Menlo"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Menlo")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Menlo")
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = DARK_BLUE
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.line_spacing = 1.0

    for current_section in doc.sections:
        header = current_section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_p.add_run("MEMSTACK PRODUCT DESIGN  |  DESKTOP AGENT CLIENT")
        set_run_font(header_run, size=8.5, color=MUTED, bold=True)

        footer = current_section.footer
        add_page_field(footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("MemStack Desktop Agent Client")
    set_run_font(run, size=27, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("统一通用 Agent 与编程 Agent 的桌面指挥中心")
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("Version", "v1.1 Draft"),
        ("Date", "2026-07-10"),
        ("Target", "macOS / Windows desktop"),
        ("Status", "Product structure and visual direction locked"),
        ("Scope", "Product PRD + UI/UX specification + selected prototype"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(12)
    rule.paragraph_format.space_after = Pt(18)
    set_paragraph_border_bottom(rule, color="2E74B5", size=12)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(12)
    run = lead.add_run(
        "核心决策：不建设两套彼此割裂的产品。Work 与 Code 共享项目、任务、线程、记忆、权限、运行状态和产物，差异只体现在任务能力集与自适应工作画布。"
    )
    set_run_font(run, size=12, color=DARK_BLUE, bold=True)

    doc.add_page_break()


def add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|`[^`]+`|\*\*[^*]+\*\*)")


def add_inline(paragraph: Paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=11, color=INK)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Menlo", size=9.2, color=DARK_BLUE)
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=11, color=INK, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=11, color=INK)


def next_abstract_num_id(numbering: CT_Numbering) -> int:
    ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    return max(ids, default=-1) + 1


def next_num_id(numbering: CT_Numbering) -> int:
    ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    return max(ids, default=0) + 1


def create_numbering(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_abstract_num_id(numbering)
    num_id = next_num_id(numbering)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_paragraph(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    paragraph._p.get_or_add_pPr().append(num_pr)
    add_inline(paragraph, text)


def add_code_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = paragraph.add_run(text if text else " ")
    set_run_font(run, name="Menlo", size=8.5, color=DARK_BLUE)


def add_image(doc: Document, alt: str, relative_path: str) -> None:
    image_path = (ROOT / relative_path).resolve()
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(6.3))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption_run = caption.add_run(f"Figure: {alt}")
    set_run_font(caption_run, size=9, color=MUTED, italic=True)


def add_markdown(  # noqa: C901, PLR0912, PLR0915
    doc: Document,
    path: Path,
    *,
    skip_title: bool = True,
) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    in_code = False
    code_language = ""
    bullet_num_id: int | None = None
    ordered_num_id: int | None = None
    skipped_first_h1 = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(segment.strip() for segment in paragraph_buffer))
        paragraph_buffer = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                in_code = False
                code_language = ""
            else:
                in_code = True
                code_language = stripped[3:].strip()
                if code_language:
                    add_code_paragraph(doc, f"[{code_language}]")
            continue
        if in_code:
            add_code_paragraph(doc, line)
            continue
        if not stripped:
            flush_paragraph()
            bullet_num_id = None
            ordered_num_id = None
            continue

        image_match = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", stripped)
        if image_match:
            flush_paragraph()
            add_image(doc, image_match.group(1), image_match.group(2))
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            if level == 1 and skip_title and not skipped_first_h1:
                skipped_first_h1 = True
                continue
            paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline(paragraph, title)
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            if bullet_num_id is None:
                bullet_num_id = create_numbering(doc, ordered=False)
            add_list_paragraph(doc, bullet_match.group(1), bullet_num_id)
            continue

        ordered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ordered_match:
            flush_paragraph()
            if ordered_num_id is None:
                ordered_num_id = create_numbering(doc, ordered=True)
            add_list_paragraph(doc, ordered_match.group(1), ordered_num_id)
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph()


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    overview = doc.add_paragraph(style="Heading 1")
    overview.add_run("Research and product framing")
    research_intro = doc.add_paragraph()
    add_inline(
        research_intro,
        "The research brief below summarizes current GitHub Copilot App and Codex / ChatGPT Desktop patterns, public UX failure signals, and the design implications for MemStack. Full source notes remain in the companion Markdown report.",
    )
    add_markdown(doc, ROOT / "01-competitive-research.md")

    doc.add_page_break()
    prd = doc.add_paragraph(style="Heading 1")
    prd.add_run("Product requirements")
    add_markdown(doc, ROOT / "02-product-prd.md")

    spec = doc.add_paragraph(style="Heading 1")
    spec.add_run("UI/UX specification")
    add_markdown(doc, ROOT / "03-ui-ux-spec.md")

    doc.core_properties.title = "MemStack Desktop Agent Client PRD"
    doc.core_properties.subject = (
        "Unified Work and Code Agent desktop product requirements and UI/UX specification"
    )
    doc.core_properties.author = "MemStack Product Design"
    doc.core_properties.keywords = "MemStack, desktop agent, Work, Code, PRD, UI UX"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
